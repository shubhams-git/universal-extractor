import os
import pathlib
import logging
from typing import Optional, Dict, Any, List

# Word document processing dependencies
try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
    logging.info("docx2txt available - Word document text extraction enabled")
except ImportError:
    DOCX2TXT_AVAILABLE = False
    logging.warning("docx2txt not available. Install: pip install docx2txt")
    docx2txt = None

try:
    from docx import Document
    PYTHON_DOCX_AVAILABLE = True
    logging.info("python-docx available - Word document structure access enabled")
except ImportError:
    PYTHON_DOCX_AVAILABLE = False
    logging.warning("python-docx not available. Install: pip install python-docx")
    Document = None

# PDF creation dependencies
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Preformatted, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
    logging.info("ReportLab available - PDF creation enabled")
except ImportError:
    REPORTLAB_AVAILABLE = False
    logging.warning("ReportLab not available. Install: pip install reportlab")
    # Create mock classes for type checking
    class _SimpleDocTemplate:
        def __init__(self, filename, **kwargs): pass
        def build(self, story): pass
    class _Paragraph:
        def __init__(self, text, style): pass
    class _Preformatted:
        def __init__(self, text, style): pass
    class _Spacer:
        def __init__(self, width, height): pass
    class _Table:
        def __init__(self, data): pass
        def setStyle(self, style): pass
    class _TableStyle:
        def __init__(self, commands): pass
    class _ParagraphStyle:
        def __init__(self, name, **kwargs): pass
    class _StyleSheet(dict):
        def add(self, style): pass
    
    SimpleDocTemplate = _SimpleDocTemplate  # type: ignore
    Paragraph = _Paragraph  # type: ignore
    Preformatted = _Preformatted  # type: ignore
    Spacer = _Spacer  # type: ignore
    Table = _Table  # type: ignore
    TableStyle = _TableStyle  # type: ignore
    ParagraphStyle = _ParagraphStyle  # type: ignore
    def getSampleStyleSheet():  # type: ignore[override]
        return _StyleSheet()
    A4 = None  # type: ignore

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class WordToPdfService:
    """
    Professional Word document to PDF conversion service.
    Handles DOCX and DOC files with structure preservation.
    """
    
    def __init__(self):
        """Initialize the Word to PDF service."""
        if not REPORTLAB_AVAILABLE:
            raise ImportError(
                "ReportLab is required for PDF creation. "
                "Install with: pip install reportlab"
            )
        
        # At least one Word processing library is required
        if not DOCX2TXT_AVAILABLE and not PYTHON_DOCX_AVAILABLE:
            raise ImportError(
                "Word processing library required. "
                "Install with: pip install docx2txt python-docx"
            )
        
        self.supported_formats = {'.docx', '.doc'}
        
        # Prefer python-docx for structure, fallback to docx2txt for text
        self.use_structured_extraction = PYTHON_DOCX_AVAILABLE
    
    def get_word_metadata(self, word_path: str) -> Dict[str, Any]:
        """Extract metadata from Word document."""
        try:
            file_stat = os.stat(word_path)
            file_size_mb = file_stat.st_size / (1024 * 1024)
            
            metadata = {
                "file_size_mb": round(file_size_mb, 3),
                "file_format": pathlib.Path(word_path).suffix.lower()
            }
            
            # Try to extract document properties using python-docx
            if PYTHON_DOCX_AVAILABLE and Document and word_path.endswith('.docx'):
                try:
                    doc = Document(word_path)
                    
                    # Core properties
                    core_props = doc.core_properties
                    metadata.update({
                        "title": core_props.title or None,
                        "author": core_props.author or None,
                        "subject": core_props.subject or None,
                        "keywords": core_props.keywords or None,
                        "created": str(core_props.created) if core_props.created else None,
                        "modified": str(core_props.modified) if core_props.modified else None,
                        "last_modified_by": core_props.last_modified_by or None
                    })
                    
                    # Document structure
                    paragraph_count = len(doc.paragraphs)
                    table_count = len(doc.tables)
                    
                    metadata.update({
                        "paragraph_count": paragraph_count,
                        "table_count": table_count,
                        "has_tables": table_count > 0
                    })
                    
                    # Word count (approximate)
                    word_count = 0
                    for paragraph in doc.paragraphs:
                        word_count += len(paragraph.text.split())
                    
                    metadata["word_count"] = word_count
                    
                except Exception as e:
                    logger.warning(f"Could not extract structured metadata: {e}")
            
            # Fallback text extraction for word count
            elif DOCX2TXT_AVAILABLE and docx2txt:
                try:
                    text = docx2txt.process(word_path)
                    metadata["word_count"] = len(text.split()) if text else 0
                except Exception as e:
                    logger.warning(f"Could not extract text for word count: {e}")
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract Word metadata: {e}")
            return {"error": str(e)}
    
    def _create_custom_styles(self):
        """Create custom styles for Word document conversion."""
        styles = getSampleStyleSheet()
        
        # Custom styles for different Word elements
        custom_styles = {
            'WordTitle': ParagraphStyle(
                'WordTitle',
                parent=styles['Title'],
                fontSize=18,
                spaceAfter=20,
                textColor=colors.darkblue if colors else 'darkblue',  # type: ignore
                alignment=1  # Center alignment
            ),
            'WordHeading1': ParagraphStyle(
                'WordHeading1',
                parent=styles['Heading1'],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=12,
                textColor=colors.darkblue if colors else 'darkblue',  # type: ignore
            ),
            'WordHeading2': ParagraphStyle(
                'WordHeading2',
                parent=styles['Heading2'],
                fontSize=12,
                spaceAfter=10,
                spaceBefore=10,
                textColor=colors.darkslategray if colors else 'darkslategray',  # type: ignore
            ),
            'WordNormal': ParagraphStyle(
                'WordNormal',
                parent=styles['Normal'],
                fontSize=10,
                leading=14,
                spaceAfter=6,
                alignment=0  # Left alignment
            ),
            'WordBullet': ParagraphStyle(
                'WordBullet',
                parent=styles['Normal'],
                fontSize=10,
                leading=14,
                spaceAfter=3,
                leftIndent=20,
                bulletIndent=10
            )
        }
        
        # Add custom styles to the existing styles
        for name, style in custom_styles.items():
            styles.add(style)
            
        return styles
    
    def convert_word_to_pdf(self, word_path: str, pdf_path: str,
                          title: Optional[str] = None,
                          preserve_structure: bool = True) -> bool:
        """
        Convert Word document to PDF with structure preservation.
        
        Args:
            word_path: Path to input Word file
            pdf_path: Path for output PDF file
            title: Custom title for the document
            preserve_structure: Whether to preserve document structure
        
        Returns:
            bool: True if conversion successful, False otherwise
        """
        try:
            logger.info(f"Converting Word to PDF: {word_path} -> {pdf_path}")
            
            # Determine conversion method
            if preserve_structure and self.use_structured_extraction and Document and word_path.endswith('.docx'):
                return self._convert_with_structure(word_path, pdf_path, title)
            else:
                return self._convert_text_only(word_path, pdf_path, title)
            
        except Exception as e:
            logger.error(f"Failed to convert Word to PDF: {e}")
            return False
    
    def _convert_with_structure(self, word_path: str, pdf_path: str, title: Optional[str] = None) -> bool:
        """Convert Word document preserving structure using python-docx."""
        try:
            doc = Document(word_path) if Document else None
            if not doc:
                return self._convert_text_only(word_path, pdf_path, title)
            
            # Create PDF document
            pdf_doc = SimpleDocTemplate(
                pdf_path,
                pagesize=A4,
                rightMargin=50,
                leftMargin=50,
                topMargin=50,
                bottomMargin=50
            )
            
            story = []
            styles = self._create_custom_styles()
            
            # Add title
            doc_title = title or (doc.core_properties.title if doc else None) or pathlib.Path(word_path).stem
            story.append(Paragraph(doc_title, styles['WordTitle']))
            story.append(Spacer(1, 20))
            
            # Process document elements
            if doc:
                for element in doc.element.body:
                    if element.tag.endswith('p'):  # Paragraph
                        para = None
                        for p in doc.paragraphs:
                            if p._element == element:
                                para = p
                                break
                        
                        if para and para.text.strip():
                            # Determine paragraph style based on Word formatting
                            style_name = self._determine_paragraph_style(para, styles)
                            story.append(Paragraph(para.text, styles[style_name]))
                    
                    elif element.tag.endswith('tbl'):  # Table - FIXED INDENTATION
                        table = None
                        for t in doc.tables:
                            if t._element == element:
                                table = t
                                break
                        
                        if table:
                            pdf_table = self._convert_word_table(table)
                            if pdf_table:
                                story.append(pdf_table)
                                story.append(Spacer(1, 12))
            
            # Build PDF
            pdf_doc.build(story)
            
            logger.info(f"Successfully converted Word to PDF with structure: {pdf_path}")
            return True
            
        except Exception as e:
            logger.error(f"Structured conversion failed: {e}")
            # Fallback to text-only conversion
            return self._convert_text_only(word_path, pdf_path, title)
    
    def _convert_text_only(self, word_path: str, pdf_path: str, title: Optional[str] = None) -> bool:
        """Convert Word document as plain text using docx2txt."""
        try:
            if not DOCX2TXT_AVAILABLE:
                logger.error("docx2txt not available for text conversion")
                return False
            
            # Extract text
            text = docx2txt.process(word_path) if docx2txt else ""
            
            if not text or not text.strip():
                logger.error(f"No text extracted from Word document: {word_path}")
                return False
            
            # Create PDF document
            pdf_doc = SimpleDocDocument(
                pdf_path,
                pagesize=A4,
                rightMargin=50,
                leftMargin=50,
                topMargin=50,
                bottomMargin=50
            )
            
            story = []
            styles = self._create_custom_styles()
            
            # Add title
            doc_title = title or pathlib.Path(word_path).stem
            story.append(Paragraph(doc_title, styles['WordTitle']))
            story.append(Spacer(1, 20))
            
            # Process text into paragraphs
            paragraphs = text.split('\n')
            
            for para_text in paragraphs:
                para_text = para_text.strip()
                if not para_text:
                    story.append(Spacer(1, 6))
                    continue
                
                # Simple style detection
                if len(para_text) < 100 and para_text.isupper():
                    # Likely a heading
                    story.append(Paragraph(para_text, styles['WordHeading1']))
                elif para_text.startswith('•') or para_text.startswith('-'):
                    # Bullet point
                    story.append(Paragraph(para_text, styles['WordBullet']))
                else:
                    # Regular paragraph
                    story.append(Paragraph(para_text, styles['WordNormal']))
            
            # Build PDF
            pdf_doc.build(story)
            
            logger.info(f"Successfully converted Word to PDF (text-only): {pdf_path}")
            return True
            
        except Exception as e:
            logger.error(f"Text-only conversion failed: {e}")
            return False
    
    def _determine_paragraph_style(self, paragraph, styles) -> str:
        """Determine appropriate PDF style for Word paragraph."""
        try:
            # Check if paragraph has a style
            if paragraph.style.name.startswith('Heading 1'):
                return 'WordHeading1'
            elif paragraph.style.name.startswith('Heading 2'):
                return 'WordHeading2'
            elif paragraph.style.name.startswith('Heading'):
                return 'WordHeading2'  # Use Heading2 for other heading levels
            elif 'Title' in paragraph.style.name:
                return 'WordTitle'
            else:
                return 'WordNormal'
        except:
            return 'WordNormal'
    
    def _convert_word_table(self, table):  # type: ignore
        """Convert Word table to ReportLab table."""
        try:
            # Extract table data
            table_data = []
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip() or " "  # Avoid empty cells
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if not table_data:
                return None
            
            # Create ReportLab table
            pdf_table = Table(table_data)
            
            # Apply basic table styling
            table_style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey if colors else 'grey'),  # type: ignore # Header background
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke if colors else 'whitesmoke'),  # type: ignore # Header text
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),  # Left align all
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),  # Header font
                ('FONTSIZE', (0, 0), (-1, 0), 10),  # Header font size
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),  # Header padding
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige if colors else 'beige'),  # type: ignore # Data background
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),  # Data font
                ('FONTSIZE', (0, 1), (-1, -1), 9),  # Data font size
                ('GRID', (0, 0), (-1, -1), 1, colors.black if colors else 'black')  # type: ignore # Grid lines
            ])
            
            pdf_table.setStyle(table_style)  # type: ignore
            return pdf_table
            
        except Exception as e:
            logger.warning(f"Failed to convert table: {e}")
            return None
    
    def batch_convert_words(self, input_files: List[str], output_dir: str = "output",
                          preserve_structure: bool = True) -> Dict[str, Any]:
        """
        Convert multiple Word files to PDF.
        
        Args:
            input_files: List of Word file paths
            output_dir: Output directory for PDF files
            preserve_structure: Whether to preserve document structure
        
        Returns:
            Dict containing conversion results
        """
        output_path = pathlib.Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)
        
        results = {
            "successful_conversions": [],
            "failed_conversions": [],
            "total_files": len(input_files),
            "success_count": 0,
            "failure_count": 0
        }
        
        for file_path in input_files:
            input_path = pathlib.Path(file_path)
            
            if not input_path.exists():
                logger.error(f"File not found: {file_path}")
                results["failed_conversions"].append({
                    "file": file_path,
                    "error": "File not found"
                })
                results["failure_count"] += 1
                continue
            
            # Check if it's a supported Word file
            if input_path.suffix.lower() not in self.supported_formats:
                logger.error(f"Unsupported Word format: {file_path}")
                results["failed_conversions"].append({
                    "file": file_path,
                    "error": "Unsupported Word format"
                })
                results["failure_count"] += 1
                continue
            
            # Generate PDF path
            pdf_filename = f"{input_path.stem}.pdf"
            pdf_path = output_path / pdf_filename
            
            # Get metadata
            metadata = self.get_word_metadata(file_path)
            
            # Convert to PDF
            success = self.convert_word_to_pdf(
                str(input_path), str(pdf_path), preserve_structure=preserve_structure
            )
            
            if success:
                results["successful_conversions"].append({
                    "input_file": file_path,
                    "output_file": str(pdf_path),
                    "metadata": metadata
                })
                results["success_count"] += 1
            else:
                results["failed_conversions"].append({
                    "file": file_path,
                    "error": "Conversion failed"
                })
                results["failure_count"] += 1
        
        return results


if __name__ == "__main__":
    """Command-line interface for Word to PDF conversion."""
    import argparse
    
    parser = argparse.ArgumentParser(description="Word to PDF Conversion Service")
    parser.add_argument("input", nargs="+", help="Input Word file(s)")
    parser.add_argument("-o", "--output", default="output", help="Output directory")
    parser.add_argument("--title", help="Custom document title")
    parser.add_argument("--no-structure", action="store_true", 
                       help="Convert as plain text (faster)")
    parser.add_argument("-v", "--verbose", action="store_true", help="Verbose logging")
    
    args = parser.parse_args()
    
    if args.verbose:
        logging.getLogger().setLevel(logging.DEBUG)
    
    try:
        service = WordToPdfService()
        
        if len(args.input) == 1:
            # Single file conversion
            input_file = args.input[0]
            input_path = pathlib.Path(input_file)
            output_path = pathlib.Path(args.output)
            output_path.mkdir(parents=True, exist_ok=True)
            
            pdf_filename = f"{input_path.stem}.pdf"
            pdf_path = output_path / pdf_filename
            
            print(f"Converting: {input_file}")
            success = service.convert_word_to_pdf(
                input_file, str(pdf_path),
                title=args.title,
                preserve_structure=not args.no_structure
            )
            
            if success:
                print(f"✅ Success: {pdf_path}")
            else:
                print("❌ Conversion failed")
        else:
            # Batch conversion
            print(f"Converting {len(args.input)} Word file(s)...")
            results = service.batch_convert_words(
                args.input, args.output, preserve_structure=not args.no_structure
            )
            
            print(f"\n📊 Results:")
            print(f"Total: {results['total_files']}")
            print(f"Success: {results['success_count']}")
            print(f"Failed: {results['failure_count']}")
            
            if results["successful_conversions"]:
                print(f"\n✅ Successful conversions:")
                for conv in results["successful_conversions"]:
                    print(f"  {conv['input_file']} -> {conv['output_file']}")
            
            if results["failed_conversions"]:
                print(f"\n❌ Failed conversions:")
                for fail in results["failed_conversions"]:
                    print(f"  {fail['file']}: {fail['error']}")
                    
    except ImportError as e:
        print(f"❌ Missing dependency: {e}")
        print("Install required packages: pip install reportlab docx2txt python-docx")
    except Exception as e:
        print(f"❌ Error: {e}")
        logger.error(f"Unexpected error: {e}", exc_info=True)